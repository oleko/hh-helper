"""
Конвертирует resume_full.md (markdown-ish текст от tailor.py) в .docx —
одна колонка, стандартные заголовки Word ("Heading 1"/"Heading 2"), без
таблиц и графики. Так резюме не теряет текст при парсинге ATS-системами
(Huntflow, Поток, E-Staff и т.п.) — они умеют читать простые .docx, но
часто ломаются на многоколоночных макетах и текстовых полях/картинках.
"""
from __future__ import annotations

import re

from docx import Document
from docx.document import Document as DocumentType

# двойные звёздочки — **bold**; одинарные — *italic* (модель использует оба:
# **компания/даты** и сразу под ней *должность* курсивом). Порядок в альтернативе
# важен — сначала пробуем "**", иначе одна "*" из "**" утащит на себя лишнее.
_INLINE_FORMAT_RE = re.compile(r"\*\*(.+?)\*\*|\*(.+?)\*")
_BOLD_LINE_RE = re.compile(r"^\*\*(.+)\*\*$")
_BULLET_RE = re.compile(r"^[-*]\s+(.*)$")

# по этим словам отличаем "заголовок раздела" (Опыт работы, Навыки, ...) от
# просто выделенной жирным строки внутри раздела (название компании и т.п.)
_SECTION_KEYWORDS = [
    "опыт работы", "профессиональный опыт", "ключевые навыки", "навыки",
    "языки", "образование", "summary", "о себе",
]


def extract_candidate_name(career_base_md: str) -> str | None:
    """Достаёт ФИО из таблицы "Быстрый снимок" карьерной базы (| ФИО | ... |)."""
    m = re.search(r"\|\s*ФИО\s*\|\s*([^|]+?)\s*\|", career_base_md)
    return m.group(1).strip() if m else None


def _is_section_header(text: str) -> bool:
    low = text.lower()
    return any(kw in low for kw in _SECTION_KEYWORDS)


def _strip_md(text: str) -> str:
    return _INLINE_FORMAT_RE.sub(lambda m: m.group(1) or m.group(2), text).strip()


def _add_runs_with_inline_bold(paragraph, text: str) -> None:
    pos = 0
    for m in _INLINE_FORMAT_RE.finditer(text):
        if m.start() > pos:
            paragraph.add_run(text[pos : m.start()])
        run = paragraph.add_run(m.group(1) if m.group(1) is not None else m.group(2))
        if m.group(1) is not None:
            run.bold = True
        else:
            run.italic = True
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


def build_resume_docx(resume_full_md: str, candidate_name: str | None = None) -> DocumentType:
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = None  # оставляем дефолт Word (11pt) — не усложняем

    if candidate_name:
        doc.add_heading(candidate_name, level=1)

    first_bold_seen = False
    for raw_line in resume_full_md.splitlines():
        line = raw_line.strip()
        if not line:
            continue

        bullet_match = _BULLET_RE.match(line)
        if bullet_match:
            doc.add_paragraph(_strip_md(bullet_match.group(1)), style="List Bullet")
            continue

        bold_match = _BOLD_LINE_RE.match(line)
        if bold_match:
            text = bold_match.group(1).strip()
            if not first_bold_seen and not candidate_name:
                # нет имени из карьерной базы — используем первую жирную строку
                # (обычно заголовок/позиционирование) как Heading 1
                doc.add_heading(text, level=1)
            elif not first_bold_seen:
                # имя уже есть заголовком — эта строка (позиционирование) идёт подзаголовком
                doc.add_heading(text, level=2)
            elif _is_section_header(text):
                doc.add_heading(text, level=2)
            else:
                p = doc.add_paragraph()
                run = p.add_run(text)
                run.bold = True
            first_bold_seen = True
            continue

        p = doc.add_paragraph()
        _add_runs_with_inline_bold(p, line)

    return doc
